import csv
import os
from docx import Document
import re

pasta = '/Users/rafaelives/Desktop/app/DemCel/agosto'

ids_procurados = ('24261','24227','24830','25053','27148','27816'
,'29775','29879','32963','33129','33828','34974'
,'35560','36231','36963','36964','36980','36994','37161','37168'
,'37178','37192','37195','37202','37204','37221','37223','37226'
,'37229','37236','37241','37244','37257','37262','37264','37272',
'37283','37286','37292','37296','37299','37302','37310','37311',
'37324','37328','37332','37335','37390','37391','37395','37429',
'37439','37452','37471','37474','37487','37489','37497','37501',
'37506','37508','37580','37581')  

resultados = {id: {'contagem': 0, 'documentos': []} for id in ids_procurados}

def extrair_texto_tabela(tabela):
    texto = ''
    for linha in tabela.rows:
        for celula in linha.cells:
            texto += celula.text + ' '
    return texto

for arquivo in os.listdir(pasta):
    if arquivo.endswith('.docx') and not arquivo.startswith('~$'):
        caminho_completo = os.path.join(pasta, arquivo)
        doc = Document(caminho_completo)
        conteudo = '\n'.join([paragrafo.text for paragrafo in doc.paragraphs])
        conteudo += '\n'.join([extrair_texto_tabela(tabela) for tabela in doc.tables])

        for id in ids_procurados:
            padrao = rf"{id}"
            qtd = len(re.findall(padrao, conteudo))
            resultados[id]['contagem'] += qtd
            if qtd > 0:
                resultados[id]['documentos'].append(arquivo)

nome_arquivo_csv = 'resultados_ids.csv'
with open(nome_arquivo_csv, 'w', newline='', encoding='utf-8') as arquivo_csv:
    escritor_csv = csv.writer(arquivo_csv)
    escritor_csv.writerow(['ID', 'Contagem Total', 'Documentos Encontrados'])
    for id, dados in resultados.items():
        escritor_csv.writerow([id, dados['contagem'], ', '.join(dados['documentos'])])

print(f"Resultados exportados para {nome_arquivo_csv}")
